import pandas as pd
import glob
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.oxml.ns import qn
import os

def loadxlsx():
    """Load the xlsx file"""
    folder_path = "./xlsx/"
    imgs_path = "./imgs/"
    xlsx_files = glob.glob(folder_path + "*.xlsx")
    

    for xlsx_file in xlsx_files: #loop through all the files
        createdocx(xlsx_file, imgs_path)

def createdocx(xlsx_file, imgs_path):
    """Create the Quickestspecs"""
    doc = Document()
    df = pd.read_excel(xlsx_file) #read the file into a pandas dataframe
    df = df[df['4RA85F [Product]'] != '##BLANK##'] #remove the ##BLANK##

    prod_name = df.loc[df['Tag'] == 'prodname', '4RA85F [Product]'].iloc[0] #get the product name

    
    ################################################################ Header
    
    header = doc.sections[0].header
    header_table = header.add_table(rows=1, cols=2, width=Inches(8))
    header_table.columns[0].width = Inches(3)
    header_table.columns[1].width = Inches(5)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_paragraph = header_table.cell(0, 0).paragraphs[0]
    header_run = header_paragraph.add_run("QuickestSpecs")
    header_run.font.size = Pt(18)
    header_run.font.bold = True
    header_paragraph = header_table.cell(0, 1).paragraphs[0]
    header_run2 = header_paragraph.add_run(prod_name)
    header_run2.font.size = Pt(14)
    header_run2.font.bold = True
    header_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    ################################################################ Footer

    footer = doc.sections[0].footer 
    footer_table = footer.add_table(rows=1, cols=3, width=Inches(8)) 

    footer_table.columns[0].width = Inches(1)
    footer_table.columns[1].width = Inches(6)
    footer_table.columns[2].width = Inches(1)

    footer_table.rows[0].height = Inches(.4)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER



    footer_paragraph = footer_table.cell(0, 0).paragraphs[0]
    footer_paragraph.add_run().add_picture(imgs_path + "hp-logo.png", width=Inches(.4), height=Inches(.4))
    footer_paragraph = footer_table.cell(0, 1).paragraphs[0]
    footer_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.text = "Worldwide — Version 1 — March 19, 2023"
    footer_paragraph = footer_table.cell(0, 2).paragraphs[0]
    footer_paragraph.text = "page x"
    




    ################################################################ Callout section

    img_path = os.path.join(imgs_path, 'c08518669.png')
    img_path2 = os.path.join(imgs_path, 'c08518762.png')

    doc.add_picture(img_path, width=Inches(6))
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Left")
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    callout_table = doc.add_table(rows=4, cols=4)
    callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(4):
        for j in range(4):
            cell = callout_table.cell(i, j)
            cell.text = str(i * 4 + j + 1)
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER


    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_picture(img_path2, width=Inches(6))

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Right")
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    callout_table2 = doc.add_table(rows=4, cols=4)
    callout_table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(4):
        for j in range(4):
            cell = callout_table2.cell(i, j)
            cell.text = str(i * 4 + j + 1)
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER


    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    ################################################################ At a glance section

    doc.add_paragraph().add_run("At a glance").bold = True

    features = df.loc[df['Tag'].str.endswith('medium'), :]
    for feature in features['4RA85F [Product]']:
        if not pd.isna(feature):
            doc.add_paragraph(feature, style='List Bullet')
    
    footnote_numbers = df[df['ContainerName'].str.endswith('(medium) Footnote Number')]
    
    for footnote in footnote_numbers['4RA85F [Product]']:
        if not pd.isna(footnote):
            doc.add_paragraph(footnote)    
    
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    ################################################################ Operating system section

    doc.add_paragraph().add_run("OPERATING SYSTEM").bold = True

    ossuppted_values = df.loc[df['Tag'] == 'ossuppted', ['ContainerName', '4RA85F [Product]']].iloc[0]
     
    ossuppted_title = doc.add_heading(ossuppted_values['ContainerName'], level=1)
    ossuppted_title.style.font.size = Pt(14)

    ossuppted_subtitle_replace = ossuppted_values['4RA85F [Product]'].replace('; ', '\n')
    ossuppted_subtitle = doc.add_heading(level=2)
    ossuppted_subtitle.add_run(ossuppted_subtitle_replace).bold = False

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


    ################################################################ Dimensions and Weight section

    doc.add_paragraph().add_run("WEIGHTS & DIMENSIONS").bold = True

    dimenus_values = df.loc[df['Tag'] == 'dimenus', ['ContainerName', '4RA85F [Product]']].iloc[0]
    dimenus_title = doc.add_heading(dimenus_values['ContainerName'], level=1)
    dimenus_title.style.font.size = Pt(14)

            
    doc.save('quickestspecs.docx')


def main():
    loadxlsx()

if __name__ == "__main__":
        main()