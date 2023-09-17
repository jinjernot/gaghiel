from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def insertHR(paragraph, thickness=12, width=Inches(8)):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def header(doc, prod_name):
    
    header = doc.sections[0].header
    
    header_table = header.add_table(rows=1, cols=2, width=Inches(8))
    header_table.columns[0].width = Inches(3)
    header_table.columns[1].width = Inches(5)

    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_paragraph = header_table.cell(0, 0).paragraphs[0]
    header_run = header_paragraph.add_run("QuickestSpecs")
    header_run.font.size = Pt(18)
    header_run.font.bold = True

    header_paragraph = header_table.cell(0, 1).paragraphs[0]
    header_run2 = header_paragraph.add_run(prod_name)
    header_run2.font.size = Pt(14)
    header_run2.font.bold = True
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    insertHR(header.add_paragraph(), thickness=30)
