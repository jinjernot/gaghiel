from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
import json

def hd_section(doc):
    doc.add_paragraph().add_run("STORAGE AND DRIVES").bold = True
    # Load data from the JSON file
    json_file_path = "json/SSD.json"  # Update with the actual path to your JSON file
    with open(json_file_path, "r") as json_file:
        data = json.load(json_file)

    # Iterate through the keys in the JSON data
    for key, attributes in data.items():
        # Create a new table for each key
        table = doc.add_table(rows=0, cols=3)
        table.allow_autofit = False  # Disable auto column width adjustment

        # Set the column widths
        table.columns[0].width = 2  # Adjust as needed
        table.columns[1].width = 2  # Adjust as needed
        table.columns[2].width = 4  # Adjust as needed

        # Set the alignment of the cells in the table
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Insert the key text in the first row of the table and make it bold
        row = table.add_row().cells
        cell = row[0]
        cell.text = key
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

        # Insert data from the JSON into the table for the current key
        for attribute, value in attributes.items():
            row = table.add_row().cells
            # Make the attribute text bold in all rows
            cell = row[1]
            cell.text = attribute
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            row[2].text = value

        # Add a page break after each table
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
