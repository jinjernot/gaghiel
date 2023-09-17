import pandas as pd
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.shared import RGBColor  # Import RGBColor

def os_section(doc, df):

    doc.add_paragraph().add_run("OPERATING SYSTEM").bold = True

    ossuppted_values = df.loc[df['Tag'] == 'ossuppted', ['ContainerName', '4RA85F [Product]']].iloc[0]
     
    ossuppted_title = doc.add_heading(ossuppted_values['ContainerName'], level=1)
    
    # Explicitly set the font size and color
    ossuppted_title.style.font.size = Pt(14)
    ossuppted_title.style.font.color.rgb = None  # Set to None to use the default color (usually black)

    ossuppted_subtitle_replace = ossuppted_values['4RA85F [Product]'].replace('; ', '\n')
    
    # Create ossuppted_subtitle as a run and set its font color to black
    ossuppted_subtitle = doc.add_paragraph().add_run(ossuppted_subtitle_replace)
    ossuppted_subtitle.bold = False
    ossuppted_subtitle.font.color.rgb = RGBColor(0, 0, 0)  # Set the font color to black

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
