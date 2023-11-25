
from quickestspects.format.hr import insertHR

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import RGBColor
from docx.shared import Pt
import pandas as pd

def system_unit_section(doc, xlsx_file, df):
    """System Unit table"""

    # Load xlst
    df = pd.read_excel(xlsx_file, sheet_name='QS-Only System Unit')

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("SYSTEM UNIT")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run().add_break()

    start_col_idx = 5
    end_col_idx = 6
    start_row_idx = 10
    end_row_idx = 52

    data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
    data_range = data_range.dropna(how='all')

    num_rows, num_cols = data_range.shape
    table = doc.add_table(rows=num_rows, cols=num_cols)

    table.alignment = WD_ALIGN_VERTICAL.CENTER

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            cell = table.cell(row_idx, col_idx)
            if not pd.isna(value):
                cell.text = str(value)
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Insert HR
    insertHR(doc.add_paragraph(), thickness=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
