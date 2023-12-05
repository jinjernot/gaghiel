
from quickestspects.blocks.paragraph import *
from quickestspects.blocks.title import *
from quickestspects.blocks.table import *
from quickestspects.format.hr import *

from docx.enum.text import WD_BREAK
import pandas as pd

def fingerprint_section(doc, xlsx_file, txt_file):
    """Fingerprint Reader QS Only Section"""

    # Load xlsx
    df = pd.read_excel(xlsx_file, sheet_name='QS-Only Fingerprint Reader')

    # Add title: Fingerprint Reader
    insertTitle(doc, "Fingerprint Reader", txt_file)

    # Add table
    insertTable(doc, df, txt_file)

    # Insert HR
    insertHR(doc.add_paragraph(), thickness=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
