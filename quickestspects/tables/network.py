
from quickestspects.blocks.paragraph import *
from quickestspects.blocks.title import *
from quickestspects.format.hr import *

from docx.enum.text import WD_BREAK
import pandas as pd

def network_section(doc, xlsx_file, txt_file):
    """Network QS Only Section"""

    # Load xlsx
    df = pd.read_excel(xlsx_file, sheet_name='QS-Only Network')

    # Add tible: Networking
    insertTitle(doc, "NETWORKING / COMMUNICATIONS", txt_file)

    for index, row in df.iterrows():
        # Check if the content in column 0 is "Table"
        if row[0] == "Table":
            # Get the starting row index for the next "Table"
            start_row_index = index + 1
            
            # Add a table with 3 columns to the Word document
            table = doc.add_table(rows=1, cols=3)
            
            # Determine the number of rows until the next "Table" is met
            end_row_index = start_row_index
            while end_row_index < len(df) and df.iloc[end_row_index, 0] != "Table":
                end_row_index += 1
            
            # Populate columns 1 and 2 with values from the DataFrame
            for i in range(start_row_index, end_row_index):
                # Populate column 1
                table.add_row().cells[1].text = str(df.iloc[i, 0])
                
                # Populate column 2
                table.rows[-1].cells[2].text = str(df.iloc[i, 1])
            
            # Insert the value next to the table into the second row of column 0
            table.cell(1, 0).text = str(row[1])  # Assuming the value is in the same row




    # Insert HR
    insertHR(doc.add_paragraph(), thickness=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
