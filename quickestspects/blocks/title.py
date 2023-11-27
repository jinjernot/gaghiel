from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def insertTitle(doc, title, txt_file):
    # Add the title to the Word document
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #paragraph.add_run().add_break()

    # Write the HTML title
    with open(txt_file, 'a') as txt:
        txt.write(f"<b><h1>{title}</h1></b>\n")


def insertSubtitle(doc, txt_file, df,  iloc_row, iloc_column):
    # Add the subtitle to the Word document
    subtitle = df.iloc[iloc_row, iloc_column]
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(subtitle)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #paragraph.add_run().add_break()

    # Write the HTML subtitle
    with open(txt_file, 'a') as txt:
        txt.write(f"<p>{subtitle}</p>\n")
