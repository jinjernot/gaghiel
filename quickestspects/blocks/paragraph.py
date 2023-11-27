from docx.enum.text import WD_BREAK
from docx.shared import RGBColor
import pandas as pd

def insertParagraph(doc, txt_file, df, iloc_row, iloc_column):
    data = df.iloc[iloc_row, iloc_column]
    paragraph = doc.add_paragraph()
    paragraph.add_run(data)

    with open(txt_file, 'a') as txt:
        txt.write(f"<p>{data}</p>\n")

def insertList(doc, txt_file, df, iloc_range, iloc_column):
    # Get the data
    items = df.iloc[iloc_range, iloc_column].tolist()
    #Remove N/A 
    items  = [data for data in items if pd.notna(data)]

    for data in items:
        run = doc.add_paragraph().add_run(data)
        #run.add_break(WD_BREAK.LINE)

        with open(txt_file, 'a') as txt:
            txt.write(f"<p>{data}</p>\n")

def insertFootnote(doc, txt_file, df, iloc_range, iloc_column):
    footnote = df.iloc[iloc_range, iloc_column].tolist()
    footnote  = [note for note in footnote if pd.notna(note)]
    paragraph = doc.add_paragraph()

    for note in footnote:
        run = paragraph.add_run(note)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.add_break(WD_BREAK.LINE)

    html_footnotes = '<div style="color: blue;">\n'
    for note in footnote:
        html_footnotes += f'  <span>{note}</span>\n'
    html_footnotes += '</div>\n'
    with open(txt_file, 'a') as txt:
            txt.write(html_footnotes)
