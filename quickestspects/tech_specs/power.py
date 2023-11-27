from quickestspects.format.hr import *
from quickestspects.blocks.title import  insertTitle

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import RGBColor
from docx.shared import Pt
import pandas as pd

def power_section(doc, txt_file, df):

    insertTitle(doc, "POWER", txt_file)


    insertHR(doc.add_paragraph(), thickness=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
