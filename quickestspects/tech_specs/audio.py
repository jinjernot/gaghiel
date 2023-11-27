from quickestspects.format.hr import *
from quickestspects.blocks.title import  insertTitle

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import RGBColor
from docx.shared import Pt
import pandas as pd

def audio_section(doc, txt_file, df):
    
    insertTitle(doc, "AUDIO / MULTIMEDIA", txt_file)

    # Add HR
    insertHR(doc.add_paragraph(), thickness=3)

    # Add HTML <hr>
    with open(txt_file, 'a') as txt:
        txt.write('<hr align="center" SIZE="2" width="100%">\n')

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
