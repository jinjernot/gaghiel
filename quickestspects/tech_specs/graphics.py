from quickestspects.blocks.paragraph import *
from quickestspects.blocks.title import *
from quickestspects.format.hr import *

from docx.enum.text import WD_BREAK


def graphics_section(doc, txt_file, df):
    """ Graphics techspecs section"""
    
    # Add the title: GRAPHICS
    insertTitle(doc, "GRAPHICS", txt_file)

    # Integrated
    insertSubtitle(doc, txt_file, df, 102, 6)
    insertList(doc, txt_file, df, slice(103, 108), 6)

    # Discrete
    insertSubtitle(doc, txt_file, df, 108, 6)
    insertList(doc, txt_file, df, slice(110, 111), 6)

    # Supports
    insertSubtitle(doc, txt_file, df, 111, 6)
    insertList(doc, txt_file, df, slice(112, 116), 6)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(117, 121), 6)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
