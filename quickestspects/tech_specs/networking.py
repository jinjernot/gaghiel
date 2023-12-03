from quickestspects.blocks.paragraph import *
from quickestspects.blocks.title import *
from quickestspects.format.hr import *

from docx.enum.text import WD_BREAK

def networking_section(doc, txt_file, df):
    """Network techspecs section"""

    # Add the title: NETWORKINGS
    insertTitle(doc, "NETWORKING", txt_file)

    # Wlan
    insertSubtitle(doc, txt_file, df, 100, 1)
    insertList(doc, txt_file, df, slice(101, 103),1)

    # Wwlan
    insertSubtitle(doc, txt_file, df, 103, 1)
    insertList(doc, txt_file, df, slice(104, 106), 1)

    # LPWAN
    insertSubtitle(doc, txt_file, df, 106, 1)
    insertList(doc, txt_file, df, slice(107, 108), 1)

    # NFC
    insertSubtitle(doc, txt_file, df, 108, 1)
    insertList(doc, txt_file, df, slice(109, 110), 1)

    # Miracast
    insertSubtitle(doc, txt_file, df, 110, 1)
    insertList(doc, txt_file, df, slice(111, 112), 1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(113, 119), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)