from quickestspects.blocks.paragraph import *
from quickestspects.blocks.title import *
from quickestspects.format.hr import *

from docx.enum.text import WD_BREAK

def storage_section(doc, txt_file, df):
    """Storage techspecs section"""

    insertTitle(doc, "STORAGE AND DRIVES", txt_file)

    # Primary Storage
    insertSubtitle(doc, txt_file, df, 67, 1)
    insertList(doc, txt_file, df, slice(68, 76), 1)

    # M2 Storage
    #insertSubtitle(doc, txt_file, df, 207, 6)
    #insertList(doc, txt_file, df, slice(208, 221), 6)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(78, 79), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)