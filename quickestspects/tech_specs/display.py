from quickestspects.blocks.paragraph import *
from quickestspects.blocks.title import *
from quickestspects.format.hr import *

from docx.enum.text import WD_BREAK

def display_section(doc, txt_file, df):
    """Display techspecs section"""

    # Add the title: DISPLAY
    insertTitle(doc, "DISPLAY", txt_file)   

    # Non-Touch
    insertSubtitle(doc, txt_file, df, 26, 1)
    insertList(doc, txt_file, df, slice(27, 31), 1)

    # Touch
    insertSubtitle(doc, txt_file, df, 31, 1)
    insertList(doc, txt_file, df, slice(32, 33), 1)

    # Display Port
    #insertSubtitle(doc, txt_file, df, 155, 1)
    #insertList(doc, txt_file, df, slice(156, 157), 1)

    # Display Support
    #insertSubtitle(doc, txt_file, df, 158, 1)
    #insertList(doc, txt_file, df, slice(159, 161), 1)

    # Display Size
    insertSubtitle(doc, txt_file, df, 33, 1)
    insertList(doc, txt_file, df, slice(34, 35), 1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(37, 41), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
