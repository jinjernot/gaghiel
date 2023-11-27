from quickestspects.blocks.paragraph import *
from quickestspects.blocks.title import *
from quickestspects.format.hr import *

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
import pandas as pd

def operating_systems_section(doc, txt_file, df):
    """Operating system techspecs section"""

    # Add the title: OPERATING SYSTEMS
    insertTitle(doc, "OPERATING SYSTEMS", txt_file)

    operating_systems = df.iloc[19:30, 6].tolist()

    operating_systems = [os for os in operating_systems if pd.notna(os)]

    total_rows = (len(operating_systems))

    os_table = doc.add_table(rows=total_rows, cols=2)
    os_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_index = 1 
    for row_index in range(total_rows):
        list_index = row_index
        if list_index < len(operating_systems):
            os_table.cell(row_index, col_index).text = str(operating_systems[list_index])
    
    preinstalled_text = df.iloc[12, 6]
    preinstalled = os_table.cell(0, 0)
    preinstalled.text = preinstalled_text
    run = preinstalled.paragraphs[0].runs[0]
    run.bold = True

    html_table = '<table class="MsoNormalTable" cellSpacing="3" cellPadding="0" width="728" border="0">\n'

    html_table += f'<tr>\n<td><strong>{preinstalled_text}</strong></td>\n</tr>\n'

    for os in operating_systems:
        html_table += f'<tr>\n<td></td>\n<td>{os}</td>\n</tr>\n'
    html_table += '</table>\n'
    with open(txt_file, 'a') as txt:
            txt.write(html_table)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(31, 36), 6)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)