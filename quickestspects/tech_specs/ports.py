from quickestspects.blocks.paragraph import *
from quickestspects.blocks.title import *
from quickestspects.format.hr import *

from docx.enum.text import WD_BREAK


def ports_section(doc, txt_file, df):
    """Ports techspecs section"""

    # Add title: PORTS
    insertTitle(doc, "PORTS", txt_file)

    # Left side
    insertSubtitle(doc, txt_file, df, 299, 1)
    insertList(doc, txt_file, df, slice(300, 305),1)

    # Right side
    insertSubtitle(doc, txt_file, df, 305, 1)
    insertList(doc, txt_file, df, slice(306, 310), 1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(311, 313), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)