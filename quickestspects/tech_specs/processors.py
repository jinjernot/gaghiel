from quickestspects.format.hr import insertHR

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor
import pandas as pd


def processors_section(doc, txt_file, df):

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("PROCESSORS")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    with open(txt_file, 'a') as txt:
        txt.write("<h1><b>PROCESSORS</h1></b>\n")

    start_col_idx = 6
    end_col_idx = 12
    start_row_idx = 52
    end_row_idx = 60

    data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
    data_range = data_range.dropna(how='all')

    num_rows, num_cols = data_range.shape
    table = doc.add_table(rows=num_rows, cols=num_cols)

    table.alignment = WD_ALIGN_VERTICAL.CENTER

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            cell = table.cell(row_idx, col_idx)

            if not pd.isna(value):
                cell.text = str(value)

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    processors_table = '<table border="1" style="border-collapse: collapse;">\n'

    for row_idx in range(data_range.shape[0]):
        processors_table += '  <tr>\n'
        for col_idx in range(data_range.shape[1]):
            value = data_range.iat[row_idx, col_idx]
            processors_table += f'    <td>{value}</td>\n' if not pd.isna(value) else '    <td></td>\n'
        processors_table += '  </tr>\n'

    processors_table += '</table>\n'

    with open(txt_file, 'a') as txt:
        txt.write(processors_table)

    run.add_break(WD_BREAK.LINE)

    processors_footnotes = df.iloc[73:80, 6].tolist()
    processors_footnotes = [os for os in processors_footnotes if pd.notna(os)]
    
    paragraph = doc.add_paragraph()

    for pro_footnote in processors_footnotes:
        run = paragraph.add_run(pro_footnote)
        run.add_break(WD_BREAK.LINE)
        run.font.color.rgb = RGBColor(0, 0, 255)
    run.add_break(WD_BREAK.LINE)

    pro_footnotes = '<div style="color: blue;">\n'

    for pro_footnote in processors_footnotes:
        pro_footnotes += f'  <span>{pro_footnote}</span>\n'

    pro_footnotes += '</div>\n'

    with open(txt_file, 'a') as txt:
            txt.write(pro_footnotes)

    insertHR(doc.add_paragraph(), thickness=3)

    with open(txt_file, 'a') as txt:
        txt.write('<hr align="center" SIZE="2" width="100%">\n')