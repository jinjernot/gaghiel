from quickestspects.blocks.paragraph import *
from quickestspects.blocks.title import *
from quickestspects.format.hr import *

from docx.enum.text import WD_BREAK

def memory_section(doc, txt_file, df):
    """Memory tectspecs section"""

    # Add the title: MEMORY
    insertTitle(doc, "MEMORY", txt_file)

    # Maximum memory
    insertSubtitle(doc, txt_file, df, 81, 1)
    insertParagraph(doc, txt_file, df, 82, 1)

    # Primary memory
    insertSubtitle(doc, txt_file, df, 83, 1)
    insertList(doc, txt_file, df, slice(84, 90), 1)

    # Memory slots
    insertSubtitle(doc, txt_file, df, 90, 1)
    insertList(doc, txt_file, df, slice(91, 93), 1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(95, 97), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)