from quickestspects.blocks.paragraph import *
from quickestspects.blocks.title import *
from quickestspects.format.hr import *

from docx.enum.text import WD_BREAK


def keyboard_section(doc, txt_file, df):

    insertTitle(doc, "KEYBOARDS / POINTING DEVICES / BUTTONS & FUNCTION KEYS", txt_file)

    # Keyboard
    insertSubtitle(doc, txt_file, df, 145, 1)
    insertList(doc, txt_file, df, slice(146, 149),1)

    # Pointing Device
    insertSubtitle(doc, txt_file, df, 149, 1)
    insertList(doc, txt_file, df, slice(150, 153), 1)

    # Function Keys
    insertSubtitle(doc, txt_file, df, 153, 1)
    insertList(doc, txt_file, df, slice(154, 168), 1)

    # Hidden Function Keys
    insertSubtitle(doc, txt_file, df, 168, 1)
    insertList(doc, txt_file, df, slice(169, 172), 1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(172, 173), 1)


    insertHR(doc.add_paragraph(), thickness=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
