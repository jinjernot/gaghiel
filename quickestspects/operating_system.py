import pandas as pd
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.shared import RGBColor  # Import RGBColor

def os_section(doc, df):

    doc.add_paragraph().add_run("OPERATING SYSTEM").bold = True

    

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
