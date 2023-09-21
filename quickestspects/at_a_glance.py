import pandas as pd
from docx.enum.text import WD_BREAK

def ataglance_section(doc, df):

    doc.add_paragraph().add_run("At a Glance").bold = True

    
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)