import re
from docx import Document

def find_text_in_square_brackets(docx_path):
    doc = Document(docx_path)

    pattern = r'\[(\d+)\]'

    matches = []

    for paragraph in doc.paragraphs:
        matches.extend(re.findall(pattern, paragraph.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches.extend(re.findall(pattern, cell.text))

    return matches

def superscript_numbers_in_square_brackets(docx_path):
    doc = Document(docx_path)

    matches = find_text_in_square_brackets(docx_path)

    for paragraph in doc.paragraphs:
        for match in matches:
            paragraph.text = re.sub(rf'\[({re.escape(match)})\]', rf'[sub]{match}[/sub]', paragraph.text, flags=re.IGNORECASE)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for match in matches:
                    cell.text = re.sub(rf'\[({re.escape(match)})\]', rf'[sub]{match}[/sub]', cell.text, flags=re.IGNORECASE)

    doc.save('output.docx')

# Example usage
docx_path = 'quickestspecs.docx'
superscript_numbers_in_square_brackets(docx_path)
