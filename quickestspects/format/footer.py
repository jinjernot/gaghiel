from quickestspects.format.hr import insertHR

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, Inches
from datetime import datetime


def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run = paragraph.add_run()
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def footer(doc, imgs_path):
    footer = doc.sections[0].footer
    insertHR(footer.add_paragraph(), thickness=10)
    footer_table = footer.add_table(rows=1, cols=3, width=Inches(8)) 

    footer_table.columns[0].width = Inches(1)
    footer_table.columns[1].width = Inches(6)
    footer_table.columns[2].width = Inches(1)

    footer_table.rows[0].height = Inches(.4)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    footer_paragraph = footer_table.cell(0, 0).paragraphs[0]
    footer_paragraph.add_run().add_picture(imgs_path + "hp-logo.png", width=Inches(.4), height=Inches(.4))
    footer_paragraph = footer_table.cell(0, 1).paragraphs[0]
    footer_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add the text before the current date
    footer_paragraph.add_run("Not all configuration components are available in all regions/countries.")
    footer_paragraph.add_run().add_break()
    for run in footer_paragraph.runs:
        run.font.size = Pt(8)

    current_date = datetime.now().strftime("%B %d, %Y")
    footer_paragraph.add_run(f"Worldwide — Version 1 — {current_date}")
    for run in footer_paragraph.runs:
        run.font.size = Pt(8)

    footer_paragraph = footer_table.cell(0, 2).paragraphs[0]
    footer_paragraph.text = "Page "
    add_page_number(footer_paragraph)
    for run in footer_paragraph.runs:
        run.font.size = Pt(8)