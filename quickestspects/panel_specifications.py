import pandas as pd
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK,WD_ALIGN_PARAGRAPH

def panel_specs(doc, df):

    globalskunumbers = df[df['Tag'] == 'globalskunumbers']['ChunkValue'].values
    if len(globalskunumbers) > 0 and not pd.isna(globalskunumbers[0]):
        doc.add_paragraph(f"Model: {globalskunumbers}")

    # Add a table with 3 columns and 20 rows
    table = doc.add_table(rows=23, cols=3)
    table.allow_autofit = False  # Disable table autofit
    
    # Set the width of each column
    col_width = Inches(2)
    for col in table.columns:
        col.width = col_width

    first_row = table.rows[0]
    first_row.cells[0].paragraphs[0].add_run("Panel Specifications")
    first_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # List of values to populate in the second column
    values_to_populate = [
        "Display Size (Diagonal)",
        "Panel Technology",
        "Max Refresh Rate",
        "Native Resolution",
        "Panel Bit Depth",
        "Aspect Ratio",
        "Brightness (Typical)",
        "Contrast Ratio (Static)",
        "Dynamic Contrast Ratio",
        "Flicker Free",
        "Pixel Pitch",
        "Pixels Per Inch (PPI)",
        "Display colors",
        "Backlight Lamp Life Minimum (To Half Brightness - In Hours)",
        "Backlight Type",
        "Screen Treatment",
        "Hardness",
        "Haze",
        "Response Time (Typical)",
        "Horizontal Viewing Angle (Typical CR>10)",
        "Vertical Viewing Angle (Typical CR>10)",
        "Panel Active Area Metric (W x H)",
        "Panel Active Area Imperial (W x H)"
    ]

    # Iterate through the rows of the table and populate the second column
    for i, value in enumerate(values_to_populate):
        row = table.rows[i]
        cell = row.cells[1]
        cell.paragraphs[0].add_run(value)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)  # Remove space after the paragraph

     # Retrieve the value for 'displaysizemet' from the DataFrame
    displaysizemet_value = df[df['Tag'] == 'displaysizemet']['ChunkValue'].values
    if len(displaysizemet_value) > 0 and not pd.isna(displaysizemet_value[0]):
        # Populate the first cell of the third column with 'displaysizemet' value
        first_cell_third_column = table.cell(0, 2)
        first_cell_third_column.paragraphs[0].add_run(displaysizemet_value[0])

    # Retrieve the value for 'displaytype' from the DataFrame
    displaytype_value = df[df['Tag'] == 'displaytype']['ChunkValue'].values
    if len(displaytype_value) > 0 and not pd.isna(displaytype_value[0]):
        # Populate the second cell of the third column with 'displaytype' value
        second_cell_third_column = table.cell(1, 2)
        second_cell_third_column.paragraphs[0].add_run(displaytype_value[0])

        # Retrieve the value for 'displayrefreshrate' from the DataFrame
    displayrefreshrate_value = df[df['Tag'] == 'displayrefreshrate']['ChunkValue'].values
    if len(displayrefreshrate_value) > 0 and not pd.isna(displayrefreshrate_value[0]):
        # Populate the third cell of the third column with 'displayrefreshrate' value
        third_cell_third_column = table.cell(2, 2)
        third_cell_third_column.paragraphs[0].add_run(displayrefreshrate_value[0])

    
     # Add the NOTE paragraph
    note_paragraph = doc.add_paragraph("NOTE: Performance specifications represent the typical specifications provided by HP's component manufacturers; actual performance may vary either higher or lower.")
    note_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
