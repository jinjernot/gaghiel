import os
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import pandas as pd


def callout_section(df, doc, imgs_path):
    img_path = os.path.join(imgs_path, 'placeholder-image.png')
    img_path2 = os.path.join(imgs_path, 'placeholder-image.png')

    doc.add_picture(img_path, width=Inches(6))
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Front")
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    callout_table = doc.add_table(rows=6, cols=2)
    callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER

 # Define the 'calloutfront_' tags you want to process
    tags_to_process = ['calloutfront_01', 'calloutfront_02', 'calloutfront_03', 'calloutfront_04', 'calloutfront_05', 'calloutfront_06']

    for row_index, tag in enumerate(tags_to_process):
        # Find the value for the current tag in the DataFrame
        value_to_populate = df[df['Tag'] == tag]['ChunkValue'].values
        if len(value_to_populate) > 0 and not pd.isna(value_to_populate[0]):
            callout_table.cell(row_index, 0).paragraphs[0].add_run(str(value_to_populate[0]))
   
   # Center the table cells
    for row in callout_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_picture(img_path2, width=Inches(6))

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Back")
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    callout_table2 = doc.add_table(rows=6, cols=2)
    callout_table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Define the 'calloutback_' tags you want to process
    tags_to_process_back = ['calloutback_01', 'calloutback_02', 'calloutback_03', 'calloutback_04', 'calloutback_05', 'calloutback_06']

    for row_index, tag in enumerate(tags_to_process_back):
        # Find the value for the current tag in the DataFrame
        value_to_populate = df[df['Tag'] == tag]['ChunkValue'].values
        if len(value_to_populate) > 0 and not pd.isna(value_to_populate[0]):
            callout_table2.cell(row_index, 0).paragraphs[0].add_run(str(value_to_populate[0]))

    # Center the table cells
    for row in callout_table2.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break() 
    section = doc.sections[-1]
    section.start_type
    section.start_type = WD_SECTION.CONTINUOUS