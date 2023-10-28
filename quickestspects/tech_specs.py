import pandas as pd
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK,WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from quickestspects.hr import insertHR

def tech_specs_section(xlsx_file, doc, df, prod_name):

    # Read data from the Excel file
    df = pd.read_excel(xlsx_file, sheet_name='Tech Specs & QS Features')

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("PRODUCT NAME")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    prod_name_paragraph = doc.add_paragraph(prod_name)
    insertHR(doc.add_paragraph(), thickness=3)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("OPERATING SYSTEMS")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    operating_systems = df.iloc[21:31, 6].tolist()

    # Filter out NaN values
    operating_systems = [os for os in operating_systems if pd.notna(os)]

    total_rows = (len(operating_systems)) # Adding 1 to round up if there's an odd number of items

    # Create the table with the dynamically determined number of rows and 2 columns
    os_table = doc.add_table(rows=total_rows, cols=2)
    os_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Populate the second column of the table with all values from operating_systems
    col_index = 1  # Select the second column
    for row_index in range(total_rows):
        list_index = row_index
        if list_index < len(operating_systems):
            os_table.cell(row_index, col_index).text = str(operating_systems[list_index])
    
    # Add "Preinstalled" to the first cell of the first column
    preinstalled = os_table.cell(0, 0)
    preinstalled.text = "Preinstalled"

    operating_systems_footnotes = df.iloc[31:36, 6].tolist()
    operating_systems_footnotes = [os for os in operating_systems_footnotes if pd.notna(os)]
    print(operating_systems_footnotes)

    # Create a new paragraph
    footnote_paragraph = doc.add_paragraph()

    # Add the data from the list to the paragraph
    for os_footnote in operating_systems_footnotes:
        run = footnote_paragraph.add_run(os_footnote)

        # Set the font color to blue
        run.font.color.rgb = RGBColor(0, 0, 255)  # RGB for blue

        run.add_break(WD_BREAK.LINE)


    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
