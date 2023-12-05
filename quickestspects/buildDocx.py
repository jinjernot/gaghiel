from quickestspects.tech_specs.tech_specs import tech_specs_section
from quickestspects.overview.overview import overview_section
from quickestspects.format.format import format_document

from quickestspects.tables.system_unit import system_unit_section
from quickestspects.tables.displays import displays_section
from quickestspects.tables.options import options_section
from quickestspects.tables.audio import audio_section
from quickestspects.tables.fingerprint import fingerprint_section
from quickestspects.tables.storage import storage_section
from quickestspects.tables.network import network_section
#from quickestspects.format.superscript import superscript

import pandas as pd
from docx import Document


def createdocx(xlsx_file, imgs_path):
    """Create the Quickestspecs"""
    
    # Variables
    doc = Document()
    txt_file = 'quickestspecs.txt'
    format_document(doc, xlsx_file, imgs_path)

    # Quickspecs sections
    #overview_section(doc, xlsx_file, txt_file, df, prod_name, imgs_path)
    tech_specs_section(doc, xlsx_file, txt_file)
    
    system_unit_section(doc, xlsx_file, txt_file)
    displays_section(doc, xlsx_file, txt_file)
    storage_section(doc, xlsx_file, txt_file)
    network_section(doc, xlsx_file, txt_file)
    audio_section(doc, xlsx_file, txt_file)
    fingerprint_section(doc, xlsx_file, txt_file)
    options_section(doc, xlsx_file, txt_file)
    #superscript(doc, xlsx_file, txt_file)

    
    docx_file = 'quickestspecs.docx'
    doc.save(docx_file)