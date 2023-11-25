from docx.enum.text import WD_BREAK,WD_ALIGN_PARAGRAPH
import pandas as pd

def ataglance_section(doc, txt_file, df):
    """At a glance"""

    # Create Subtitle
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("At a Glance")
    run.bold = True

    with open(txt_file, 'a') as txt:
        txt.write("<h1><b>At a Glance</h1></b>\n")

    with open(txt_file, 'a') as txt:
        txt.write("<tbody>\n")
        txt.write("<tr style='HEIGHT: 210.5pt'>\n")
        txt.write("<td style='HEIGHT: 210.5pt; WIDTH: 537.05pt; PADDING-BOTTOM: 0.75pt; PADDING-TOP: 0.75pt; PADDING-LEFT: 0.75pt; PADDING-RIGHT: 0.75pt' vAlign='top' width='716'>\n")

    # Add at a glance content
    at_a_glance = df.iloc[70:85, 6].tolist()
    for item in at_a_glance:
        if pd.notna(item):
            doc.add_paragraph(item, style='List Bullet')

            #HTML
            with open(txt_file, 'a') as txt:
                txt.write(f"<p>{item}</p>\n")
    with open(txt_file, 'a') as txt:
        txt.write("</td></tr>\n")

    # Add at a glance note
    note_paragraph = doc.add_paragraph("NOTE: See important legal disclosures for all listed specs in their respective features sections")
    note_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in note_paragraph.runs:
        run.bold = True
    with open(txt_file, 'a') as txt:
        txt.write("<b>NOTE: See important legal disclosures for all listed specs in their respective features sections</b>\n")
    with open(txt_file, 'a') as txt:
        txt.write('<hr align="center" SIZE="2" width="100%">\n')
    with open(txt_file, 'a') as txt:
        txt.write("<p class='MsoNormal' style='LINE-HEIGHT: 115%'></table>\n")

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)