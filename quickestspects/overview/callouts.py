from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Pt, Inches
import pandas as pd
import os


def callout_section(doc, txt_file, df, prod_name, imgs_path):
    prodname_paragraph = doc.add_paragraph()
    run = prodname_paragraph.add_run(prod_name)
    run.font.size = Pt(12)
    run.bold = True
    prodname_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    with open(txt_file, 'a') as txt:
        txt.write(f"<html><head><title>{prod_name}</title>\n")
        txt.write('<meta content="text/html; charset=utf-8" http-equiv="Content-Type">\n')
        txt.write('<meta name="Generator" content="Microsoft Word 15 (filtered)">\n')
        txt.write("</head>\n")
        txt.write("<h1>Overview</h1>\n")
        txt.write(f"<p style='font-size:14pt;'><strong>{prod_name}</strong></p>\n")


    img_path = os.path.join(imgs_path, 'image001.png')
    img_path2 = os.path.join(imgs_path, 'image002.png')

    img_html_code = f'<img src="{img_path}" alt="Product Image" width="702" height="561">'
    img_html_code2 = f'<img src="{img_path2}" alt="Product Image" width="702" height="561">'

    with open(txt_file, 'a') as txt:
        txt.write(img_html_code + '\n')
        txt.write("<p>Left</p>\n")

    doc.add_picture(img_path, width=Inches(6))

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Front")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    callouts = df.iloc[11:31, 6].tolist()

    callouts = [tag for tag in callouts if pd.notna(tag)]

    total_rows = (len(callouts) + 1) // 2

    callout_table = doc.add_table(rows=total_rows, cols=2)
    callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index in range(total_rows):
        for col_index in range(2):
            list_index = row_index * 2 + col_index
            if list_index < len(callouts):
                callout_table.cell(row_index, col_index).text = str(callouts[list_index])

    for row in callout_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    table_html = '<table border="1" style="border-collapse: collapse;">\n'
    for i in range(0, len(callouts), 2):
        row_html = f'<tr>\n<td>{callouts[i]}</td>\n'
        if i + 1 < len(callouts):
            row_html += f'<td>{callouts[i + 1]}</td>\n'
        else:
            row_html += '<td></td>\n'
        row_html += '</tr>\n'
        table_html += row_html

    table_html += '</table>\n'

    with open(txt_file, 'a') as txt:
        txt.write(table_html)

    with open(txt_file, 'a') as txt:
        txt.write('<hr align="center" SIZE="2" width="100%">\n')

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_picture(img_path2, width=Inches(6))

    with open(txt_file, 'a') as txt:
        txt.write(img_html_code2 + '\n')
        txt.write("<p>Right</p>\n")

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Back")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tags_to_process_back = df.iloc[40:60, 6].tolist()

    filtered_tags2 = [tag for tag in tags_to_process_back if pd.notna(tag)]

    total_rows = (len(filtered_tags2) + 1) // 2 

    callout_table2 = doc.add_table(rows=total_rows, cols=2)
    callout_table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index in range(total_rows):
        for col_index in range(2):
            list_index = row_index * 2 + col_index
            if list_index < len(filtered_tags2):
                callout_table2.cell(row_index, col_index).text = str(filtered_tags2[list_index])

    for row in callout_table2.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)

    table_html2 = '<table border="1" style="border-collapse: collapse;">\n'
    for i in range(0, len(filtered_tags2), 2):
        row_html = f'<tr>\n<td>{filtered_tags2[i]}</td>\n'
        if i + 1 < len(filtered_tags2):
            row_html += f'<td>{filtered_tags2[i + 1]}</td>\n'
        else:
            row_html += '<td></td>\n'
        row_html += '</tr>\n'
        table_html2 += row_html

    table_html2 += '</table>\n'

    with open(txt_file, 'a') as txt:
        txt.write(table_html2)

    with open(txt_file, 'a') as txt:
        txt.write('<hr align="center" SIZE="2" width="100%">\n')

    doc.add_page_break()
    section = doc.sections[-1]
    section.start_type
    section.start_type = WD_SECTION.CONTINUOUS
