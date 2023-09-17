from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING
from quickestspects.hr import insertHR

def header(doc, prod_name):
    
    header = doc.sections[0].header
    
    header_table = header.add_table(rows=1, cols=2, width=Inches(8))
    header_table.columns[0].width = Inches(3)
    header_table.columns[1].width = Inches(5)

    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left-aligned text
    header_paragraph = header_table.cell(0, 0).paragraphs[0]
    header_run = header_paragraph.add_run("QuickestSpecs")
    header_run.font.size = Pt(24)
    header_run.font.bold = True

    # Right-aligned text
    header_paragraph = header_table.cell(0, 1).paragraphs[0]
    header_run2 = header_paragraph.add_run(prod_name)
    header_run2.font.size = Pt(14)
    header_run2.font.bold = True
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    insertHR(header.add_paragraph(), thickness=30)