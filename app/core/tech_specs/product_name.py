from app.core.format.hr import *

from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import pandas as pd

def product_name_section(doc, html_file, prod_name):
    """Product name section"""

    paragraph = doc.add_paragraph()
    run = paragraph.add_run("PRODUCT NAME")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    with open(html_file, 'a', encoding='utf-8') as txt:
        txt.write("<h1><b>Technical Specifications</h1></b>\n")
        txt.write("<h1><b>PRODUCT NAME</h1></b>\n")
        txt.write(f"<p>{prod_name}</p>\n")
    
    paragraph = doc.add_paragraph(prod_name)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(html_file)