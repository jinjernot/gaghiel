from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK

def power_section(doc, txt_file, df):
    """Power techspecs section"""

    # Add the title: POWER
    insertTitle(doc, "POWER", txt_file)

    # Power Supply
    insertSubtitle(doc, txt_file, df, 263, 1)
    insertList(doc, txt_file, df, slice(264, 267),1)

    # Battery
    insertSubtitle(doc, txt_file, df, 267, 1)
    insertList(doc, txt_file, df, slice(268, 270), 1)

    # Battery Recharge Time
    insertSubtitle(doc, txt_file, df, 270, 1)
    insertList(doc, txt_file, df, slice(271, 272), 1)

    # Power Cord
    insertSubtitle(doc, txt_file, df, 272, 1)
    insertList(doc, txt_file, df, slice(273, 274), 1)

    # Battery life
    insertSubtitle(doc, txt_file, df, 274, 1)
    insertList(doc, txt_file, df, slice(275, 279), 1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(281, 285), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
