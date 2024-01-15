from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK

def dimensions_section(doc, html_file, df):
    """Dimensions techspecs section"""

    # Add title
    insertTitle(doc, "WEIGHTS & DIMENSIONS", html_file)

    # Product Weight
    insertSubtitle(doc, html_file, df, 288, 1)
    insertList(doc, html_file, df, slice(289, 290),1)

    # Product Dimensions (w x d x h)
    insertSubtitle(doc, html_file, df, 290, 1)
    insertList(doc, html_file, df, slice(291, 292), 1)

    # Package Dimensions (w x d x h)
    insertSubtitle(doc, html_file, df, 293, 1)
    insertList(doc, html_file, df, slice(294, 295), 1)

    # Footnotes
    insertFootnote(doc, html_file, df, slice(297, 299), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(html_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)