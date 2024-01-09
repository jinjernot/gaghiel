from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK


def ports_section(doc, txt_file, df):
    """Ports techspecs section"""

    # Add title: PORTS
    insertTitle(doc, "PORTS", txt_file)

    # Left side
    insertSubtitle(doc, txt_file, df, 301, 1)
    insertList(doc, txt_file, df, slice(302, 306),1)

    # Right side
    insertSubtitle(doc, txt_file, df, 306, 1)
    insertList(doc, txt_file, df, slice(307, 311), 1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(312, 315), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)