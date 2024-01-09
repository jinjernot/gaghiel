from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK


def service_section(doc, txt_file, df):
    """Service and support techspecs section"""

    # Add the title: SERVICE AND SUPPORT
    insertTitle(doc, "SERVICE AND SUPPORT", txt_file)

    # Service and Support
    insertParagraph(doc, txt_file, df, 317, 1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(320, 321), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
