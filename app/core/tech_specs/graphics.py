from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK

def graphics_section(doc, txt_file, df):
    """Graphics techspecs section"""
    
    # Add the title: GRAPHICS
    insertTitle(doc, "GRAPHICS", txt_file)

    # Integrated
    insertSubtitle(doc, txt_file, df, 13, 1)
    insertList(doc, txt_file, df, slice(14, 16), 1)

    # Discrete
    insertSubtitle(doc, txt_file, df, 16, 1)
    insertList(doc, txt_file, df, slice(17, 18), 1)

    # Supports
    insertSubtitle(doc, txt_file, df, 18, 1)
    insertList(doc, txt_file, df, slice(19, 21), 1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(22, 24), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
