from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK

def memory_section(doc, html_file, df):
    """Memory tectspecs section"""

    # Add the title: MEMORY
    insertTitle(doc, "MEMORY", html_file)

    # Maximum memory
    insertSubtitle(doc, html_file, df, 81, 1)
    insertParagraph(doc, html_file, df, 82, 1)

    # Primary memory
    insertSubtitle(doc, html_file, df, 83, 1)
    insertList(doc, html_file, df, slice(84, 90), 1)

    # Memory slots
    insertSubtitle(doc, html_file, df, 90, 1)
    insertList(doc, html_file, df, slice(91, 93), 1)

    # Footnotes
    insertFootnote(doc, html_file, df, slice(95, 97), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(html_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)