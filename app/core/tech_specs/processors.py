from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
import pandas as pd

def processors_section(doc, txt_file, df):
    """Processors techspecs section"""

    # Add the title: PROCESSORS
    insertTitle(doc, "PROCESSORS", txt_file)

    start_col_idx = 6
    end_col_idx = 12
    start_row_idx = 52
    end_row_idx = 60

    data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
    data_range = data_range.dropna(how='all')

    num_rows, num_cols = data_range.shape
    table = doc.add_table(rows=num_rows, cols=num_cols)

    table.alignment = WD_ALIGN_VERTICAL.CENTER

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            cell = table.cell(row_idx, col_idx)

            if not pd.isna(value):
                cell.text = str(value)

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    processors_table = '<table border="1" style="border-collapse: collapse;">\n'

    for row_idx in range(data_range.shape[0]):
        processors_table += '  <tr>\n'
        for col_idx in range(data_range.shape[1]):
            value = data_range.iat[row_idx, col_idx]
            processors_table += f'    <td>{value}</td>\n' if not pd.isna(value) else '    <td></td>\n'
        processors_table += '  </tr>\n'

    processors_table += '</table>\n'

    with open(txt_file, 'a') as txt:
        txt.write(processors_table)

    run.add_break(WD_BREAK.LINE)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(73, 80), 6)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)