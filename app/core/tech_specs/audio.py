from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK

def audio_section(doc, txt_file, df):
    
    insertTitle(doc, "AUDIO / MULTIMEDIA", txt_file)

    # Audio
    insertSubtitle(doc, txt_file, df, 121, 1)
    insertList(doc, txt_file, df, slice(122, 126),1)

    # Speaker Power
    insertSubtitle(doc, txt_file, df, 127, 1)
    insertList(doc, txt_file, df, slice(128, 129),1)

    # Camera
    insertSubtitle(doc, txt_file, df, 129, 1)
    insertList(doc, txt_file, df, slice(130, 132),1)

    # Sensors
    insertSubtitle(doc, txt_file, df, 132, 1)
    insertList(doc, txt_file, df, slice(133, 138),1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(140, 142), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
