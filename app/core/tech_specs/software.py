from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK

def software_section(doc, txt_file, df):
    """Software and security techspecs section"""

    # Add the title: SOFTWARE AND SECURITY
    insertTitle(doc, "SOFTWARE AND SECURITY", txt_file)

    # Software
    insertSubtitle(doc, txt_file, df, 173, 1)
    insertList(doc, txt_file, df, slice(175, 191),1)

    # Manageability Features
    insertSubtitle(doc, txt_file, df, 191, 1)
    insertList(doc, txt_file, df, slice(193, 200), 1)

    # Security Management
    insertSubtitle(doc, txt_file, df, 203, 1)
    insertList(doc, txt_file, df, slice(202, 212), 1)

    # Security- TPM
    insertSubtitle(doc, txt_file, df, 212, 1)
    insertList(doc, txt_file, df, slice(214, 223), 1)

    # TCG TPM 2.0
    #insertSubtitle(doc, txt_file, df, 215, 1)
    #insertList(doc, txt_file, df, slice(216, 217), 1)

    # FIPS 140-2 Compliant: Yes
    #insertSubtitle(doc, txt_file, df, 217, 1)
    #insertList(doc, txt_file, df, slice(219, 221), 1)

    # BIOS
    insertSubtitle(doc, txt_file, df, 224, 1)
    insertList(doc, txt_file, df, slice(226, 232), 1)

    # Smartcard Reader
    insertSubtitle(doc, txt_file, df, 232, 1)
    insertList(doc, txt_file, df, slice(233, 235), 1)

    # IPv6 Support
    insertSubtitle(doc, txt_file, df, 236, 1)
    insertList(doc, txt_file, df, slice(237, 238), 1)

    # FirstNet Certified
    insertSubtitle(doc, txt_file, df, 238, 1)
    insertList(doc, txt_file, df, slice(239, 243), 1)

    # Footnotes
    insertFootnote(doc, txt_file, df, slice(245, 261), 1)


    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(txt_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
