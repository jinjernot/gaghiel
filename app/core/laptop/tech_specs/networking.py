from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK

def networking_section(doc, html_file, df):
    """Network techspecs section"""

    # Add the title: NETWORKINGS
    #insertTitle(doc, "NETWORKING", html_file)

    insertList(doc, html_file, df, "Networking /Communications")

    # Wlan
    #insertSubtitle(doc, html_file, df, 100, 1)
    #insertList(doc, html_file, df, slice(101, 103),1)

    # Wwlan
    #insertSubtitle(doc, html_file, df, 103, 1)
    #insertList(doc, html_file, df, slice(104, 106), 1)

    # LPWAN
    #insertSubtitle(doc, html_file, df, 106, 1)
    #insertList(doc, html_file, df, slice(107, 108), 1)

    # NFC
    #insertSubtitle(doc, html_file, df, 108, 1)
    #insertList(doc, html_file, df, slice(109, 110), 1)

    # Miracast
    #insertSubtitle(doc, html_file, df, 110, 1)
    #insertList(doc, html_file, df, slice(111, 112), 1)

    # Footnotes
    #insertFootnote(doc, html_file, df, slice(114, 119), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(html_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)