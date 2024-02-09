from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK


def ports_section(doc, html_file, df):
    """Ports techspecs section"""

    # Add title: PORTS
    #insertTitle(doc, "PORTS", html_file)
    insertList(doc, html_file, df, "Ports/Slots")

    # Left side
    #insertSubtitle(doc, html_file, df, 301, 1)
    #insertList(doc, html_file, df, slice(302, 306),1)

    # Right side
   #insertSubtitle(doc, html_file, df, 306, 1)
    #insertList(doc, html_file, df, slice(307, 311), 1)

    # Footnotes
    #insertFootnote(doc, html_file, df, slice(313, 315), 1)

    # HR
    #insertHR(doc.add_paragraph(), thickness=3)
    #insertHTMLhr(html_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)