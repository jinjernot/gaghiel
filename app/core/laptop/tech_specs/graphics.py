from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK

def graphics_section(doc, html_file, df):
    """Graphics techspecs section"""
    
    insertList(doc, html_file, df, "Graphics")

    # List of words to be bolded
    bold_words = ["Integrated", "Supports"]

    # Iterate through the paragraphs in the document
    for paragraph in doc.paragraphs:
        # Iterate through the runs in each paragraph
        for run in paragraph.runs:
            # Iterate through the bold words list
            for word in bold_words:
                # If the word is found in the run text, make it bold
                if word.lower() in run.text.lower():
                    run.bold = True


    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(html_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
