from app.core.format.hr import *

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import pandas as pd

def product_name_section(doc, file):
    """Product name section"""

    try:

        # Read Excel file
        df = pd.read_excel(file.stream, sheet_name='Callouts', engine='openpyxl')
        prod_name = df.columns[1]

        # Add title in Word document
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("Product Name")
        run.font.size = Pt(12)
        run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add product name to Word document
        paragraph = doc.add_paragraph(prod_name)

        # Add horizontal line
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)