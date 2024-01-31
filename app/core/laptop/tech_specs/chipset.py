from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK

def chipset_section(doc, html_file, df):

    # Add the title
    insertTitle(doc, "CHIPSET", html_file)
    
    # Add paragraph
    insertParagraph(doc, html_file, df, 90, 6)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(html_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
