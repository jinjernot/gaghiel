
from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK
import pandas as pd
import docx

def table_column_widths(table, widths):
    """Set the column widths for a table."""
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def processors_section(doc, file, html_file):
    """Processors techspecs section"""

    try:
        # Read Excel file
        df = pd.read_excel(file.stream, sheet_name='QS-Only Processors', engine='openpyxl')

        # Add title
        insert_title(doc, "Processors", html_file)

        # Dynamically fetch the column names from the third row (index 2) that have data
        third_row = df.iloc[3]  # Selecting the third row (index 2, but 1-based)

        # Filter to keep columns that have data in the third row
        filtered_df = df.loc[:, ~third_row.isna()]

        # Remove the first row
        filtered_df = filtered_df.iloc[3:]

        # Replace "NaN" string values with an empty string
        filtered_df = filtered_df.fillna('')

        # Convert filtered dataframe to a list of lists (data) for the table
        data = filtered_df.values.tolist()

        # Add the data as a table to the document
        table = doc.add_table(rows=1, cols=len(data[0]))
        table.autofit = True

        # Add table data
        for row in data:
            # Check if the row is empty
            if not any(row):
                break  # Exit the loop if the row is empty
            
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)
                
                # Bold the text if the cell contains "Processor Family"
                if str(cell) == "Processor Family":
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

        # Remove the first row
        if len(table.rows) > 1:  # Ensure there are rows to delete
            table.rows[0]._element.getparent().remove(table.rows[0]._element)

        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
        # Set the column widths of the table
        table_column_widths(table, [docx.shared.Inches(2)] * len(data[0]))  # Set each column width to 2 inches

        doc.add_paragraph()

        # After adding the table, continue processing the DataFrame
        footnotes_index = df[df.eq('Footnotes').any(axis=1)].index.tolist()
        if footnotes_index:
            footnotes_index = footnotes_index[0]  # Assuming there's only one "Footnotes" row
            footnotes_data = df.iloc[footnotes_index + 1:]  # Get data after "Footnotes" row
            footnotes_data = footnotes_data.dropna(how='all')  # Drop rows with all NaN values
            
            # Iterate over rows of footnotes_data and add them to the document
            for _, row in footnotes_data.iterrows():
                row_values = row.dropna().tolist()
                if row_values:
                    # Add row values as a paragraph with specified font color
                    paragraph = doc.add_paragraph(" - ".join(map(str, row_values)))
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 153)  # Set font color to blue    

        # HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)
        insert_html_horizontal_line(html_file)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)