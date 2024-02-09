from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK


def keyboard_section(doc, html_file, df):

    #insertTitle(doc, "KEYBOARDS / POINTING DEVICES / BUTTONS & FUNCTION KEYS", html_file)

    insertList(doc, html_file, df, "Keyboards/Pointing Devices/Buttons & Function Keys")
    # Keyboard
    #insertSubtitle(doc, html_file, df, 144, 1)
    #insertList(doc, html_file, df, slice(145, 148),1)

    # Pointing Device
    #insertSubtitle(doc, html_file, df, 148, 1)
    #insertList(doc, html_file, df, slice(149, 151), 1)

    # Function Keys
    #insertSubtitle(doc, html_file, df, 151, 1)
    #insertList(doc, html_file, df, slice(153, 165), 1)

    # Hidden Function Keys
    #insertSubtitle(doc, html_file, df, 166, 1)
    #insertParagraph(doc, html_file, df, 167, 1)


    # Footnotes
    #insertFootnote(doc, html_file, df, slice(170, 171), 1)


    insertHR(doc.add_paragraph(), thickness=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
