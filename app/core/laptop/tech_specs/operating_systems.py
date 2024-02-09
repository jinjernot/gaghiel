from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
import pandas as pd

def operating_systems_section(doc, html_file, df):
    """Operating system techspecs section"""

    # Add the title: OPERATING SYSTEMS
    #insertTitle(doc, "OPERATING SYSTEMS", html_file)
    
    insertList(doc, html_file, df, "Operating Systems")

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(html_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)