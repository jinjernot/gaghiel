from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK

def storage_section(doc, html_file, df):
    """Storage techspecs section"""

    insertTitle(doc, "STORAGE AND DRIVES", html_file)

    insertList(doc, html_file, df, "Primary Storage")
    # Primary Storage
    #insertSubtitle(doc, html_file, df, 67, 1)
    #insertList(doc, html_file, df, slice(68, 76), 1)

    # M2 Storage
    #insertSubtitle(doc, html_file, df, 207, 6)
    #insertList(doc, html_file, df, slice(208, 221), 6)

    # Footnotes
    #insertFootnote(doc, html_file, df, slice(78, 79), 1)

    # HR
    insertHR(doc.add_paragraph(), thickness=3)
    insertHTMLhr(html_file)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)