from app.core.format.table import table_column_widths
from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK
from docx.shared import Inches

def table_column_widths(table, widths):
    """Set the column widths for a table."""
    for idx, width in enumerate(widths):
        table.columns[idx].width = width

def docking_section(doc, df):
    """Docking Table"""

    try:
        # Add title: DOCKING
        insert_title(doc, "Docking (Sold Separately)")

        for index, row in df.iterrows():
            # Check if the content in column 0 is "Docking"
            if row[1] == "Docking (sold separately)":
                # Add a table with 2 columns to the Word document
                table = doc.add_table(rows=1, cols=2)

                # Set the column widths
                table_column_widths(table, [Inches(3), Inches(5)])

                # Populate columns 0 and 1 with values from the DataFrame
                for i in range(index + 1, len(df)):
                    if df.iloc[i, 0] == "Container Name":
                        break  # Exit the loop when encountering the next "Table"
                    else:
                        # Populate column 0 and set text to bold
                        cell_0 = table.add_row().cells[0]
                        cell_0.text = str(df.iloc[i, 0])
                        for paragraph in cell_0.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

                        # Populate column 1 and set text to bold
                        cell_1 = table.rows[-1].cells[1]
                        cell_1.text = str(df.iloc[i, 1])

                # Remove the first row from the table
                table.rows[0]._element.getparent().remove(table.rows[0]._element)
                
                # Add a paragraph break after the table
                doc.add_paragraph()

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)
        
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)