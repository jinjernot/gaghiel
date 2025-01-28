from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.blocks.table import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK
import pandas as pd

def network_section(doc, file):
    """Network QS Only Section"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only Network', engine='openpyxl')
        
        # Replace "_x000D_" with an empty string in the DataFrame
        df = df.replace("_x000D_", "", regex=True)

        # Add title: Networking
        insert_title(doc, "NETWORKING / COMMUNICATION")

        # Add table
        insert_table(doc, df)

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)