from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.blocks.table import *
from app.core.format.hr import *
import pandas as pd

from docx.enum.text import WD_BREAK
from docx.shared import Inches

def fingerprint_section(doc, file, html_file):
    """Fingerprint Reader QS Only Section"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only Fingerprint Reader', engine='openpyxl')

        # Add title: Fingerprint Reader
        insert_title(doc, "FINGERPRINT READER", html_file)

        start_col_idx = 0
        end_col_idx = 1
        start_row_idx = 3
        end_row_idx = 30

        data_range = df.iloc[start_row_idx:end_row_idx + 1, start_col_idx:end_col_idx + 1]
        data_range = data_range.dropna(how='all')

        num_rows, num_cols = data_range.shape
        table = doc.add_table(rows=num_rows, cols=num_cols)

        # Define column widths
        column_widths = (Inches(3), Inches(5))

        # Set column widths
        table_column_widths(table, column_widths)

        for row_idx in range(num_rows):
            for col_idx in range(num_cols):
                value = data_range.iat[row_idx, col_idx]
                cell = table.cell(row_idx, col_idx)
                if not pd.isna(value):
                    cell.text = str(value)
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        html_table = '<table class="MsoNormalTable" cellSpacing="3" cellPadding="0" width="728" border="0">\n'
        for row_idx in range(num_rows):
            html_table += "  <tr>\n"
            for col_idx in range(num_cols):
                value = data_range.iat[row_idx, col_idx]
                html_table += f"    <td>{value}</td>\n"
            html_table += "  </tr>\n"
        html_table += "</table>"

        with open(html_file, 'a', encoding='utf-8') as txt:
            txt.write(html_table)

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)
    
def table_column_widths(table, widths):
    """Set the column widths for a table."""
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
