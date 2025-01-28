from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.blocks.table import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK
import pandas as pd

def power_section(doc, file):
    """Power QS Only Section"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only Power', engine='openpyxl')

        # Add title: Power
        insert_title(doc, "POWER")

        # Add table
        insert_table(doc, df)

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)