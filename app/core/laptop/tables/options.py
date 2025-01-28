from app.core.format.table import table_column_widths
from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *
from docx.shared import Inches

from docx.enum.text import WD_BREAK
import pandas as pd

def options_section(doc, file):
    """Options QS Only Section"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only Options', engine='openpyxl')

        # Add title: Options
        insert_title(doc, "OPTIONS")

        start_col_idx = 0
        end_col_idx = 2
        start_row_idx = 0
        end_row_idx = 299

        data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
        data_range = data_range.dropna(how='all')

        num_rows, num_cols = data_range.shape
        table = doc.add_table(rows=num_rows+1, cols=num_cols)  # Adding 1 for the header row

        table_column_widths(table, (Inches(2), Inches(4), Inches(2)))

        # Adding table headers as the first row
        for col_idx in range(num_cols):
            header = df.columns[col_idx]
            cell = table.cell(0, col_idx)
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        # Populating table cells with data
        for row_idx in range(1, num_rows+1):  # Start from the second row for data
            for col_idx in range(num_cols):
                value = data_range.iat[row_idx-1, col_idx]  # Adjust row index
                cell = table.cell(row_idx, col_idx)
                if not pd.isna(value):
                    cell.text = str(value)

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)