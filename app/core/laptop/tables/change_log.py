from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *
from app.core.format.table import table_column_widths
from docx.enum.text import WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
import pandas as pd

def change_log_section(doc, file, html_file):
    """Changelog table"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='Changelog', engine='openpyxl')

        # Add title: Changelog
        insert_title(doc, "CHANGELOG", html_file)

        start_col_idx = 1
        end_col_idx = 4
        start_row_idx = 4
        end_row_idx = 12

        data_range = df.iloc[start_row_idx:end_row_idx + 1, start_col_idx:end_col_idx + 1]
        data_range = data_range.dropna(how='all')

        num_rows, num_cols = data_range.shape
        table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')

        table_column_widths(table, (Inches(2), Inches(2), Inches(2), Inches(2)))

        table.alignment = WD_ALIGN_VERTICAL.CENTER

        for row_idx in range(num_rows):
            for col_idx in range(num_cols):
                value = data_range.iat[row_idx, col_idx]
                cell = table.cell(row_idx, col_idx)
                if not pd.isna(value):
                    cell.text = str(value)

                    # Bold the text in the first row
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True

        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
        insert_paragraph(doc, html_file, df, 15, 1)
        insert_paragraph(doc, html_file, df, 16, 1)
        # Insert HR

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)
# Define the function insert_paragraph here if it's not already defined
