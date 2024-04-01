from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.blocks.table import *
from app.core.format.hr import *
import pandas as pd

from docx.enum.text import WD_BREAK

def displays_section(doc, file, html_file):
    """Displays QS Only Section"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only Displays', engine='openpyxl')

        # Add title: Displays
        insert_title(doc, "DISPLAYS", html_file)

        # Add table
        insert_table(doc, df, html_file)

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)