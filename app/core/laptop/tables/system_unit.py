from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.format.hr import *
from docx.shared import Inches

from docx.enum.text import WD_BREAK
import pandas as pd

def system_unit_section(doc, file):
    """System Unit table"""

    try:
        # Load xlsx
        df = pd.read_excel(file.stream, sheet_name='QS-Only System Unit', engine='openpyxl')

        # Add title: SYSTEM UNIT
        insert_title(doc, "SYSTEM UNIT")

        start_col_idx = 0
        end_col_idx = 1
        start_row_idx = 2
        end_row_idx = 41

        data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
        data_range = data_range.dropna(how='all')

        num_rows, num_cols = data_range.shape
        table = doc.add_table(rows=num_rows, cols=num_cols)

        # Define column widths
        column_widths = (Inches(3), Inches(5))

        # Set column widths
        table_column_widths(table, column_widths)

        for row_idx in range(num_rows):
            for col_idx in range(num_cols):
                value = data_range.iat[row_idx, col_idx]
                cell = table.cell(row_idx, col_idx)
                if not pd.isna(value):
                    cell.text = str(value)

        # Bold the first column
        for row in table.rows:
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                               
        doc.add_paragraph()

        # After adding the table, continue processing the DataFrame
        footnotes_index = df[df.eq('Footnotes').any(axis=1)].index.tolist()
        if footnotes_index:
            footnotes_index = footnotes_index[0]  # Assuming there's only one "Footnotes" row
            footnotes_data = df.iloc[footnotes_index + 1:]  # Get data after "Footnotes" row
            footnotes_data = footnotes_data.dropna(how='all')  # Drop rows with all NaN values
            
            # Iterate over rows of footnotes_data and add them to the document
            for _, row in footnotes_data.iterrows():
                row_values = row.dropna().tolist()
                if row_values:
                    # Add row values as a paragraph with specified font color
                    paragraph = doc.add_paragraph(" - ".join(map(str, row_values)))
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 153)  # Set font color to blue    

        # Insert HR
        insert_horizontal_line(doc.add_paragraph(), thickness=3)

        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    except Exception as e:
        print(f"An error occurred: {e}")
        return str(e)

def table_column_widths(table, widths):
    """Set the column widths for a table."""
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
