from docx.enum.text import WD_BREAK,WD_ALIGN_PARAGRAPH
from app.core.blocks.title import *
import pandas as pd

def ataglance_section(doc, df):
    """At a glance section"""

    insert_title(doc, "At a glance ")
    
    # Add at a glance content
    at_a_glance = df.iloc[70:85, 6].tolist()
    for item in at_a_glance:
        if pd.notna(item):
            doc.add_paragraph(item, style='List Bullet')
            
    # Add at a glance note
    note_paragraph = doc.add_paragraph("NOTE: See important legal disclosures for all listed specs in their respective features sections")
    note_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in note_paragraph.runs:
        run.bold = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)