from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from app.core.format.table import table_column_widths
from docx.shared import Pt, Inches
import pandas as pd
import requests
from io import BytesIO
from app.core.format.hr import *
import os

def download_image(url):
    """Download image from URL and return the image data."""
    response = requests.get(url)
    if response.status_code == 200:
        return BytesIO(response.content)
    else:
        return None

def get_temp_filename(counter, suffix=".png"):
    """Generate a fixed temporary file name with a three-digit counter."""
    return f"image{counter:03d}{suffix}"


def callout_section(doc, file, prod_name, df):
    """Add Callout Section"""

    # Add the product name
    prodname_paragraph = doc.add_paragraph()
    run = prodname_paragraph.add_run(prod_name)
    run.font.size = Pt(14)
    run.bold = True
    prodname_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Read text from the external file
    with open("/home/garciagi/qs/app/core/format/styles.txt", 'r', encoding='utf-8') as external_txt:
    #with open("app/core/format/styles.txt", 'r', encoding='utf-8') as external_txt:
        external_text = external_txt.read()

    # Read the doc
    #df = pd.read_excel(file, sheet_name='Callouts')
    df = pd.read_excel(file.stream, sheet_name='Callouts', engine='openpyxl')

    # Set the target directory
    target_directory = '/home/garciagi/qs'
    #target_directory = '.'

    # Get image URLs from the DataFrame
    img_url1 = df.iloc[4, 0]  # Assuming column 0, row 5
    img_url2 = df.iloc[11, 0]  # Assuming column 0, row 12

    # Initialize image counter
    img_counter = 1

    # Download images
    img_data1 = download_image(img_url1)
    img_data2 = download_image(img_url2)

    # Generate fixed temporary file names
    img_filename1 = get_temp_filename(img_counter)
    img_filename2 = get_temp_filename(img_counter + 1)

    # Increment image counter for the next iteration
    img_counter += 2

    # Save images to the specified directory
    img_filepath1 = os.path.join(target_directory, img_filename1)
    img_filepath2 = os.path.join(target_directory, img_filename2)

    with open(img_filepath1, "wb") as img_file1:
        img_file1.write(img_data1.getvalue())

    with open(img_filepath2, "wb") as img_file2:
        img_file2.write(img_data2.getvalue())

    paragraph_with_image = doc.add_paragraph()
    run = paragraph_with_image.add_run()
    run.add_picture(os.path.join(target_directory, img_filename1), width=Inches(5))
    
    # Center the paragraph
    paragraph_with_image.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Front subtitle
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Front")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Table "Front"
    start_col_idx = 1
    end_col_idx = 4
    start_row_idx = 4
    end_row_idx = 9

    data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
    data_range = data_range.dropna(how='all')

    num_rows, num_cols = data_range.shape
    table = doc.add_table(rows=num_rows, cols=num_cols)

    table_column_widths(table, (Inches(.5), Inches(3.5), Inches(.5), Inches(3.5)))

    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            cell = table.cell(row_idx, col_idx)
            if not pd.isna(value):
                if isinstance(value, (int, float)):
                    cell.text = str(int(value))
                else:
                    cell.text = str(value)
                        # Insert HR
    insert_horizontal_line(doc.add_paragraph(), thickness=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    paragraph_with_image = doc.add_paragraph()
    run = paragraph_with_image.add_run()
    run.add_picture(os.path.join(target_directory, img_filename2), width=Inches(5))

    # Center the paragraph
    paragraph_with_image.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Right subtitle
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Sides")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # add table 'Sides'
    start_col_idx = 1
    end_col_idx = 4
    start_row_idx = 11
    end_row_idx = 22

    data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
    data_range = data_range.dropna(how='all')

    num_rows, num_cols = data_range.shape
    table = doc.add_table(rows=num_rows, cols=num_cols)

    table_column_widths(table, (Inches(.5), Inches(3.5), Inches(.5), Inches(3.5)))

    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            cell = table.cell(row_idx, col_idx)
            if not pd.isna(value):
                if isinstance(value, (int, float)):
                    cell.text = str(int(value))
                else:
                    cell.text = str(value)

    # Insert HR
    insert_horizontal_line(doc.add_paragraph(), thickness=3)

    doc.add_page_break()
    section = doc.sections[-1]
    section.start_type
    section.start_type = WD_SECTION.CONTINUOUS
