
from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.blocks.table import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK
import pandas as pd

def network_section(doc, file, html_file):
    """Network QS Only Section"""

    # Load xlsx
    df = pd.read_excel(file, sheet_name='QS-Only Network')

    # Add title: Networking
    insertTitle(doc, "NETWORKING / COMMUNICATIONS", html_file)

    # Add table
    insertTable(doc, df, html_file)

    # Insert HR
    insertHR(doc.add_paragraph(), thickness=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
