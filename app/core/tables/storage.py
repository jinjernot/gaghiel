
from app.core.blocks.paragraph import *
from app.core.blocks.title import *
from app.core.blocks.table import *
from app.core.format.hr import *

from docx.enum.text import WD_BREAK
import pandas as pd

def storage_section(doc, file, txt_file):
    """Storage QS Only Section"""

    # Load xlsx
    df = pd.read_excel(file, sheet_name='QS-Only Storage')

    # Add title: Storage Reader
    insertTitle(doc, "STORAGE", txt_file)

    # Add table
    insertTable(doc, df, txt_file)

    # Insert HR
    insertHR(doc.add_paragraph(), thickness=3)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
