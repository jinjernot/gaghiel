from app.core.format.header import header
from app.core.format.footer import footer
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import json

from docx.shared import Pt

def read_bold_words_from_json(json_file):
    """
    Read bold words from a JSON file.

    Parameters:
        json_file (str): The path to the JSON file.

    Returns:
        list: A list of bold words.
    """
    with open(json_file, 'r') as f:
        data = json.load(f)
        return data.get('bold_words', [])

def set_margins(doc):
    """
    Set document margins.

    Parameters:
        doc (docx.Document): The Word document object.
    """
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(20)
        section.right_margin = Pt(20)
        section.top_margin = Pt(20)
        section.bottom_margin = Pt(20)

def set_default_font(doc):
    """
    Set default font for the document.

    Parameters:
        doc (docx.Document): The Word document object.
    """
    styles = doc.styles
    default_style = styles['Normal']
    font = default_style.font
    font.name = 'HP Simplified'
    font.size = Pt(10)

def apply_bold_font(doc, bold_words):
    """
    Apply bold font to specific words in the document.

    Parameters:
        doc (docx.Document): The Word document object.
        bold_words (list): A list of words to be bolded.
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for word in bold_words:
                if word in run.text:
                    index = run.text.find(word)
                    if index != -1 and (index == 0 or run.text[index - 1] == '\n'):
                        run.bold = True

def format_document(doc, file, imgs_path):
    """
    Apply formatting to the document.

    Parameters:
        doc (docx.Document): The Word document object.
        file (str): The path to the Word document.
        imgs_path (str): The path to the images directory.
    """
    bold_words = read_bold_words_from_json('/home/garciagi/qs/app/core/format/bold_words.json')
    header(doc, file)
    footer(doc, imgs_path)
    set_margins(doc)
    set_default_font(doc)
    apply_bold_font(doc, bold_words)

    # Apply cell spacing to all tables
    for table in doc.tables:
        table.style.paragraph_format.space_before = Pt(0)
        table.style.paragraph_format.space_after = Pt(0)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style.paragraph_format.space_before = Pt(0)
                    paragraph.style.paragraph_format.space_after = Pt(0)
