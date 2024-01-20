from app.core.format.hr import *

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

import pandas as pd

def header(doc,file):
    """Generate the Word Header"""

    # Get product name
    #df = pd.read_excel(file,sheet_name = 'Tech Specs & QS Features') 
    df = pd.read_excel(file.stream, sheet_name='Tech Specs & QS Features', engine='openpyxl')

    prod_name = df.columns[1]
    header = doc.sections[0].header
    
    header_table = header.add_table(rows=1, cols=2, width=Inches(8))
    header_table.columns[0].width = Inches(4)
    header_table.columns[1].width = Inches(4)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_paragraph = header_table.cell(0, 0).paragraphs[0]
    header_run = header_paragraph.add_run("QuickestSpecs")
    header_run.font.size = Pt(27)
    header_run.font.bold = True

    header_paragraph = header_table.cell(0, 1).paragraphs[0]
    header_run2 = header_paragraph.add_run(prod_name)
    header_run2.font.size = Pt(12)
    header_run2.font.bold = True
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    insertHR(header.add_paragraph(), thickness=30)