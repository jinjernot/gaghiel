from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from app.core.format.table import table_column_widths
from docx.shared import Pt, Inches
import pandas as pd
import requests
from io import BytesIO
import hashlib
import os

def download_image(url):
    """Download image from URL and return the image data."""
    response = requests.get(url)
    if response.status_code == 200:
        return BytesIO(response.content)
    else:
        return None

def get_temp_filename(url, suffix=".png"):
    """Generate a unique temporary file name based on the URL."""
    hash_object = hashlib.md5(url.encode())
    return f"{hash_object.hexdigest()}{suffix}"


def callout_section(doc, file, html_file, prod_name, imgs_path, df):
    """Add Callout Section"""

    # Add the product name
    prodname_paragraph = doc.add_paragraph()
    run = prodname_paragraph.add_run(prod_name)
    run.font.size = Pt(12)
    run.bold = True
    prodname_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # add HTML Headers
    with open(html_file, 'a', encoding='utf-8') as txt:
        txt.write(f"<html><head><title>{prod_name}</title>\n")
        txt.write('<meta content="text/html; charset=utf-8" http-equiv="Content-Type">\n')
        txt.write('<meta name="Generator" content="Microsoft Word 15 (filtered)">\n')
        txt.write("</head>\n")

        #Add title and product name
        txt.write("<h1><b>Overview</h1></b>\n")
        txt.write(f"<p style='font-size:14pt;'><strong>{prod_name}</strong></p>\n")

    # Assuming 'file' is the path to your Excel file
    df = pd.read_excel(file, sheet_name='Callouts')

    # Get image URLs from the DataFrame
    img_url1 = df.iloc[4, 0]  # Assuming column 0, row 5
    img_url2 = df.iloc[11, 0]  # Assuming column 0, row 12
    print (img_url1, img_url2)

    # Download images
    img_data1 = download_image(img_url1)
    img_data2 = download_image(img_url2)

    # Generate unique temporary file names
    img_filename1 = get_temp_filename(img_url1)
    img_filename2 = get_temp_filename(img_url2)

    # Save images to temporary files
    with open(img_filename1, "wb") as img_file1:
        img_file1.write(img_data1.getvalue())

    with open(img_filename2, "wb") as img_file2:
        img_file2.write(img_data2.getvalue())

    # Insert images into the Word document
    doc.add_picture(img_filename1, width=Inches(6))
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Image HTML Tags
    img_html_code = f'<img src="{img_url1}" alt="Product Image" width="702" height="561">'
    img_html_code2 = f'<img src="{img_url2}" alt="Product Image" width="702" height="561">'

    # Left image HTML subtitle
    with open(html_file, 'a', encoding='utf-8') as txt:
        txt.write(img_html_code + '\n')
        txt.write("<b><p>Front</p></b>\n")

    # Add Front subtitle
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Front")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add Table "Front"
    start_col_idx = 1
    end_col_idx = 4
    start_row_idx = 4
    end_row_idx = 9

    data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
    data_range = data_range.dropna(how='all')

    num_rows, num_cols = data_range.shape
    table = doc.add_table(rows=num_rows, cols=num_cols)

    table_column_widths(table, (Inches(.5), Inches(3.5), Inches(.5), Inches(3.5)))

    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            cell = table.cell(row_idx, col_idx)
            if not pd.isna(value):
                if isinstance(value, (int, float)):
                    cell.text = str(int(value))
                else:
                    cell.text = str(value)

    html_table = '<table class="MsoNormalTable" cellSpacing="3" cellPadding="0" width="728" border="0">\n'
    for row_idx in range(num_rows):
        html_table += "  <tr>\n"
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            html_table += f"    <td>{value}</td>\n"
        html_table += "  </tr>\n"
    html_table += "</table>"

    with open(html_file, 'a', encoding='utf-8') as txt:
        txt.write(html_table)

    # add HTML <hr>
    with open(html_file, 'a', encoding='utf-8') as txt:
        txt.write('<hr align="center" SIZE="2" width="100%">\n')

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # add docx image
    doc.add_picture(img_filename2, width=Inches(6))

    # Right image HTML subtitle
    with open(html_file, 'a', encoding='utf-8') as txt:
        txt.write(img_html_code2 + '\n')
        txt.write("<b><p>Sides</p></b>\n")

    # Add Right subtitle
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Sides")
    run.font.size = Pt(12)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # add table 'Sides'
    start_col_idx = 1
    end_col_idx = 4
    start_row_idx = 11
    end_row_idx = 22

    data_range = df.iloc[start_row_idx:end_row_idx+1, start_col_idx:end_col_idx+1]
    data_range = data_range.dropna(how='all')

    num_rows, num_cols = data_range.shape
    table = doc.add_table(rows=num_rows, cols=num_cols)

    table_column_widths(table, (Inches(.5), Inches(3.5), Inches(.5), Inches(3.5)))

    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx in range(num_rows):
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            cell = table.cell(row_idx, col_idx)
            if not pd.isna(value):
                if isinstance(value, (int, float)):
                    cell.text = str(int(value))
                else:
                    cell.text = str(value)

    html_table = '<table class="MsoNormalTable" cellSpacing="3" cellPadding="0" width="728" border="0">\n'
    for row_idx in range(num_rows):
        html_table += "  <tr>\n"
        for col_idx in range(num_cols):
            value = data_range.iat[row_idx, col_idx]
            html_table += f"    <td>{value}</td>\n"
        html_table += "  </tr>\n"
    html_table += "</table>"

    with open(html_file, 'a', encoding='utf-8') as txt:
        txt.write(html_table)
    # add HTML <hr>
    with open(html_file, 'a', encoding='utf-8') as txt:
        txt.write('<hr align="center" SIZE="2" width="100%">\n')

    doc.add_page_break()
    section = doc.sections[-1]
    section.start_type
    section.start_type = WD_SECTION.CONTINUOUS
