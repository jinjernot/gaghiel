from docx.shared import Pt

def insert_title(doc, title):
    """
    Insert a title into the Word document.

    Parameters:
        doc (docx.Document): The Word document object.
        title (str): The title text.
    """
    # Add the title to the Word document
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.font.size = Pt(12)
    run.bold = True

def insert_subtitle(doc, df, iloc_row, iloc_column):
    """
    Insert a subtitle into the Word document.

    Parameters:
        doc (docx.Document): The Word document object.
        df (pandas.DataFrame): The DataFrame containing subtitle data.
        iloc_row (int): The row index in the DataFrame.
        iloc_column (int): The column index in the DataFrame.
    """
    # Add the subtitle to the Word document
    paragraph = doc.add_paragraph()
    subtitle = df.iloc[iloc_row, iloc_column]
    run = paragraph.add_run(subtitle)
    run.bold = True