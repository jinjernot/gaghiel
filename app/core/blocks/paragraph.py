from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.shared import RGBColor
from app.core.format.hr import *

def insert_paragraph(doc, df, iloc_row, iloc_column):
    """
    Insert a paragraph into both the Word document.

    Parameters:
        doc (docx.Document): The Word document object.
        df (pandas.DataFrame): The DataFrame containing the data.
        iloc_row (int): The row index in the DataFrame.
        iloc_column (int): The column index in the DataFrame.
    """
    data = df.iloc[iloc_row, iloc_column]
    paragraph = doc.add_paragraph()
    paragraph.add_run(data)

def process_footnotes(doc, footnotes):
    """
    Add footnotes to the Word document with blue font color.

    Parameters:
        doc (docx.Document): The Word document object.
        footnotes (list): The list of footnotes to be added.
    """
    if not footnotes:
        return

    paragraph = doc.add_paragraph()
    for index, data in enumerate(footnotes):
        run = paragraph.add_run(data)
        run.font.color.rgb = RGBColor(0, 0, 153)
        
        if index < len(footnotes) - 1:
            run.add_break(WD_BREAK.LINE)

def insert_list(doc, df, start_value):
    """
    Insert a list into the Word document.

    Parameters:
        doc (docx.Document): The Word document object.
        df (pandas.DataFrame): The DataFrame containing the data.
        start_value (str): The starting value for the list.
    """
    if start_value not in df.iloc[:, 1].tolist():
        print(f"Error: '{start_value}' not found in DataFrame.")
        return
    
    start_index = df.index[df.iloc[:, 1] == start_value].tolist()[0]
    next_value_indices = df.iloc[start_index:, 1][df.iloc[start_index:, 1] == 'Value'].index.tolist()
    
    if not next_value_indices:
        print("Error: 'Value' not found after", start_value)
        return
    next_value_index = next_value_indices[0]
    
    items = df.iloc[start_index:next_value_index, 1].tolist()
    
    # Identify footnotes based on the content starting with '['
    footnotes_index = next(
        (i for i, item in enumerate(items) if item.startswith("[")),
        len(items)  # Default to end of list if no footnotes are found
    )
    non_footnotes = items[:footnotes_index]
    footnotes = items[footnotes_index:]
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(start_value.upper()) 
    paragraph = doc.add_paragraph()
    run.font.size = Pt(12)
    run.bold = True
    run.add_break(WD_BREAK.LINE)

    for index, data in enumerate(non_footnotes[1:], start=1):
        run = paragraph.add_run(data)
        
        if index < len(non_footnotes) - 1:
            run.add_break(WD_BREAK.LINE)
    
    run.add_break(WD_BREAK.LINE)
    process_footnotes(doc, footnotes)

    insert_horizontal_line(doc.add_paragraph(), thickness=3)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    
def insert_footnote(doc, df, iloc_range, iloc_column):
    """
    Insert a footnote into both the Word document.

    Parameters:
        doc (docx.Document): The Word document object.
        df (pandas.DataFrame): The DataFrame containing the data.
        iloc_range (slice): The slice range for selecting footnotes.
        iloc_column (int): The column index in the DataFrame.
    """
    footnote = df.iloc[iloc_range, iloc_column].tolist()

    paragraph = doc.add_paragraph()

    for index, note in enumerate(footnote):
        run = paragraph.add_run(note)
        run.font.color.rgb = RGBColor(0, 0, 153)
        
        if index < len(footnote) - 1:
            run.add_break(WD_BREAK.LINE)