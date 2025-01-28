from app.core.format.hr import *
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor
from docx.shared import Inches


def process_footnotes(doc, footnotes):
    """
    Process footnotes and add them to the Word document with blue font color.

    Parameters:
        doc (docx.Document): The Word document object.
        footnotes (list): The list of footnotes to be added.
    """
    paragraph = doc.add_paragraph()
    for index, data in enumerate(footnotes):
        # Skip footnotes containing unwanted values
        if "Container Name" in data or "Wireless WAN" in data:
            continue
        run = paragraph.add_run(data)
        run.font.color.rgb = RGBColor(0, 0, 153)
        
        if index < len(footnotes) - 1 and footnotes[index + 1].strip():  # Check if the next footnote is not empty
            run.add_break(WD_BREAK.LINE)

def insert_table(doc, df):
    """
    Insert tables into the Word document.

    Parameters:
        doc (docx.Document): The Word document object.
        df (pandas.DataFrame): The DataFrame containing the table data.
    """
    footnotes = []  # To store footnotes temporarily
    
    # Remove NaN values and empty rows from the DataFrame
    df.fillna('', inplace=True)
    df.dropna(how='all', inplace=True)
    
    for index, row in df.iterrows():
        if row[0] == "Table":
            # Calculate page width
            page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
            
            # Add a table to the Word document
            table = doc.add_table(rows=1, cols=3)
            column_widths = (Inches(2), Inches(2), Inches(4))
            for column, width in zip(table.columns, column_widths):
                column.width = width
            for i in range(index + 1, len(df)):
                if df.iloc[i, 0] == "Table":
                    break
                elif df.iloc[i, 0] == "Footnotes":
                    footnotes = []
                    for j in range(i + 1, len(df)):
                        if df.iloc[j, 0] == "Table":
                            break
                        footnotes.append(str(df.iloc[j, 0]))
                    break
                else:
                    cell_1 = table.add_row().cells[1]
                    cell_1.text = str(df.iloc[i, 0])
                    for paragraph in cell_1.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                    cell_2 = table.rows[-1].cells[2]
                    cell_2.text = str(df.iloc[i, 1])

            cell_0 = table.cell(1, 0)
            cell_0.text = str(row[1])
            for paragraph in cell_0.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            table.rows[0]._element.getparent().remove(table.rows[0]._element)
            
            if footnotes:
                process_footnotes(doc, footnotes)
                footnotes = []
            doc.add_paragraph()