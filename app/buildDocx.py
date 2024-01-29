from app.core.tech_specs.tech_specs import tech_specs_section
from app.core.overview.overview import overview_section
from app.core.format.format import format_document
from app.core.tables.system_unit import system_unit_section
from app.core.tables.displays import displays_section
from app.core.tables.audio import audio_section
from app.core.tables.fingerprint import fingerprint_section
from app.core.tables.storage import storage_section
from app.core.tables.network import network_section
from app.core.tables.options import options_section
from app.core.tables.change_log import change_log_section

import pandas as pd
from docx import Document
from zipfile import ZipFile
from docx2pdf import convert


imgs_path = "/home/garciagi/qs/imgs/"
#imgs_path = "./imgs/"

def createdocx(file):
    """Create the quickspecs"""
    
    # Variables
    doc = Document()
    html_file = '/home/garciagi/qs/quickspecs.html'
    #html_file = 'quickspecs.html'

    format_document(doc, file, imgs_path)

    # Quickspecs sections
    overview_section(doc, file, html_file)
    tech_specs_section(doc, file, html_file)
    system_unit_section(doc, file, html_file)
    displays_section(doc, file, html_file)
    storage_section(doc, file, html_file)
    network_section(doc, file, html_file)
    audio_section(doc, file, html_file)
    fingerprint_section(doc, file, html_file)
    options_section(doc, file, html_file)
    change_log_section(doc, file, html_file)

    docx_file = '/home/garciagi/qs/quickspecs.docx'
    #docx_file = 'quickspecs.docx'

    doc.save(docx_file)

    # Convert DOCX to PDF using docx2pdf
    #convert(docx_file)

    # Create a zip file and add specific files to it
    zip_file_name = '/home/garciagi/qs/quickspecs.zip'
    #zip_file_name = 'quickspecs.zip'

    with ZipFile(zip_file_name, 'w') as zipf:
        zipf.write(html_file, arcname='quickspecs.html')
        zipf.write(docx_file, arcname='quickspecs.docx')
        zipf.write("/home/garciagi/qs/image001.png", arcname='image001.png')
        zipf.write("/home/garciagi/qs/image002.png", arcname='image002.png')
     #   zipf.write(docx_file.replace('.docx', '.pdf'))
        
