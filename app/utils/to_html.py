import docx
import os

def extract_images_from_docx(doc):
    images = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype and "/media/" in rel.target_ref:
            image_part = rel.target_part
            image_data = image_part.blob
            images[image_part.partname.split("/")[-1]] = image_data
    return images

def save_images(images, output_folder):
    os.makedirs(output_folder, exist_ok=True)
    for image_name, image_data in images.items():
        image_path = os.path.join(output_folder, image_name)
        with open(image_path, 'wb') as image_file:
            image_file.write(image_data)

def convert_word_to_html(word_file, output_html, image_output_folder):
    # Load the Word document
    doc = docx.Document(word_file)

    # Extract images
    images = extract_images_from_docx(doc)

    # Save images to a folder
    save_images(images, image_output_folder)

    # Create an HTML template
    html_template = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <title>Word to HTML Conversion</title>
    </head>
    <body>
        {content}
    </body>
    </html>
    """

    # Process paragraphs
    paragraphs_html = ""
    for paragraph in doc.paragraphs:
        paragraphs_html += f"<p>{paragraph.text}</p>"

    # Process tables
    tables_html = ""
    for table in doc.tables:
        tables_html += "<table>"
        for row in table.rows:
            tables_html += "<tr>"
            for cell in row.cells:
                tables_html += f"<td>{cell.text}</td>"
            tables_html += "</tr>"
        tables_html += "</table>"

    # Process images
    images_html = ""
    for image_name in images.keys():
        images_html += f"<img src='{os.path.join(image_output_folder, image_name)}' alt='Image'>"

    # Combine paragraphs, tables, and images in the HTML content
    html_content = paragraphs_html + tables_html + images_html

    # Insert the HTML content into the template
    final_html = html_template.format(content=html_content)

    # Save the HTML content to the output file
    with open(output_html, 'w', encoding='utf-8') as html_file:
        html_file.write(final_html)

    print(f"HTML content saved to {output_html}")

if __name__ == "__main__":
    word_file_path = "quickspecs.docx"
    output_html_path = "output.html"
    image_output_folder = "images"
    convert_word_to_html(word_file_path, output_html_path, image_output_folder)
